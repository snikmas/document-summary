import streamlit as st
import time
import json
import io
import requests
from fpdf import FPDF
from docx import Document

st.set_page_config(
    page_title="DocPipeline",
    page_icon=":material/description:",
    layout="centered"
)

API_URL = 'http://localhost:8000'
SUPPORTED_TYPES = ['pdf', 'docx', 'txt', 'html', 'csv']

# 1. init state
if 'job_id' not in st.session_state:
    st.session_state.job_id = None
if 'poll_start' not in st.session_state:
    st.session_state.poll_start = None

POLL_TIMEOUT = 300  # 5 minutes max polling


# 2. ui
st.title("Document Intelligence Pipeline")
st.caption("Upload any document — PDF, DOCX, TXT, HTML, or CSV — and get an instant AI-powered summary.")

# U4: restrict uploader to supported types
uploaded_file = st.file_uploader("Upload your file", type=SUPPORTED_TYPES)

# reset job when file is removed from the uploader
if uploaded_file is None and st.session_state.job_id is not None:
    st.session_state.job_id = None
    st.session_state.poll_start = None
    st.rerun()

# U8: file info preview after selection
if uploaded_file is not None and st.session_state.job_id is None:
    file_size = uploaded_file.size
    if file_size < 1024:
        size_str = f"{file_size} B"
    elif file_size < 1024 * 1024:
        size_str = f"{file_size / 1024:.1f} KB"
    else:
        size_str = f"{file_size / (1024 * 1024):.1f} MB"

    ext = uploaded_file.name.rsplit('.', 1)[-1].upper() if '.' in uploaded_file.name else 'Unknown'
    st.info(f"**{uploaded_file.name}** — {ext} file, {size_str}")

# 3. handle upload
if uploaded_file is not None and st.session_state.job_id is None:
    if st.button("Summarize it", type="primary"):
        filename = uploaded_file.name
        file_bytes = uploaded_file.read()
        content_type = uploaded_file.type

        try:
            with st.spinner("Uploading file..."):
                res = requests.post(
                    f'{API_URL}/process',
                    files={'file': (filename, file_bytes, content_type)}
                )
                res.raise_for_status()
            st.session_state.job_id = res.json()['job_id']
            st.session_state.poll_start = time.time()
            st.rerun()
        except requests.exceptions.ConnectionError:
            st.error("Cannot reach the server. Make sure the backend is running.")
        except requests.exceptions.RequestException as e:
            st.error(f"Upload failed: {e}")


# 4. poll if job is running
if st.session_state.job_id:
    try:
        response = requests.get(f'{API_URL}/jobs/{st.session_state.job_id}')
        response.raise_for_status()
        status_data = response.json()
        status = status_data['status']
    except requests.exceptions.RequestException:
        st.error("Cannot reach the server. Make sure the backend is running.")
        st.stop()

    st.divider()

    if status == 'done':
        st.success("Analysis complete!")

        try:
            resp = requests.get(f'{API_URL}/jobs/{st.session_state.job_id}/result')
            resp.raise_for_status()
            result = resp.json()
        except requests.exceptions.RequestException:
            st.error("Failed to fetch results. Please try again.")
            st.stop()

        col1, col2 = st.columns(2)
        col1.metric("Language", result.get('language', 'N/A'))
        col2.metric("Word Count", result.get('word_count', 'N/A'))

        st.subheader("Summary")
        st.markdown(result['summary'])

        st.subheader("Key Points")
        points_md = '\n'.join(f"- {point}" for point in result['key_points'])
        st.markdown(points_md)

        st.divider()
        dl_col, _, reset_col = st.columns([1, 2, 1])
        with dl_col:
            with st.popover(":material/download: Download"):
                # Markdown
                md_text = f"# Summary\n\n{result['summary']}\n\n## Key Points\n\n"
                md_text += '\n'.join(f"- {p}" for p in result['key_points'])
                md_text += f"\n\n**Language:** {result.get('language', 'N/A')}  \n"
                md_text += f"**Word Count:** {result.get('word_count', 'N/A')}\n"
                st.download_button(
                    label=":material/description: Markdown (.md)",
                    data=md_text,
                    file_name="summary.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

                # PDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(text="Summary", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
                pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(w=0, text=result['summary'])
                pdf.ln(6)
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(text="Key Points", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
                pdf.set_font("Helvetica", size=11)
                for p in result['key_points']:
                    pdf.multi_cell(w=0, text=f"  \u2022  {p}")
                    pdf.ln(2)
                pdf.ln(4)
                pdf.set_font("Helvetica", "I", 10)
                pdf.cell(text=f"Language: {result.get('language', 'N/A')}  |  Word Count: {result.get('word_count', 'N/A')}")
                st.download_button(
                    label=":material/picture_as_pdf: PDF (.pdf)",
                    data=bytes(pdf.output()),
                    file_name="summary.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )

                # DOCX
                doc = Document()
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(result['summary'])
                doc.add_heading("Key Points", level=1)
                for p in result['key_points']:
                    doc.add_paragraph(p, style="List Bullet")
                doc.add_paragraph(f"Language: {result.get('language', 'N/A')}  |  Word Count: {result.get('word_count', 'N/A')}")
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button(
                    label=":material/draft: Word (.docx)",
                    data=buf.getvalue(),
                    file_name="summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        with reset_col:
            if st.button("Process another file"):
                st.session_state.job_id = None
                st.session_state.poll_start = None
                st.rerun()

    elif status in ('processing', 'pending'):
        elapsed = time.time() - (st.session_state.poll_start or time.time())
        if elapsed > POLL_TIMEOUT:
            st.error("Processing is taking too long. The server may be overloaded or the AI service is unavailable.")
            if st.button("Try again"):
                st.session_state.job_id = None
                st.session_state.poll_start = None
                st.rerun()
        else:
            if status == 'pending':
                msg = "Queued, waiting to start..."
            else:
                msg = "Analyzing your document..."
            with st.spinner(msg):
                time.sleep(2)
            st.rerun()

    elif status == 'failed':
        error_detail = status_data.get('error', '')
        if error_detail:
            st.error(f"Processing failed: {error_detail}")
        else:
            st.error("Processing failed. The file may be unsupported or corrupted.")

        if st.button("Try again"):
            st.session_state.job_id = None
            st.session_state.poll_start = None
            st.rerun()
