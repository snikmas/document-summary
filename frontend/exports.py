"""Pure, reusable summary export functions."""

from __future__ import annotations

import io
from collections.abc import Mapping
from typing import Any

from docx import Document
from fpdf import FPDF


class PdfCharacterError(ValueError):
    """Raised when the lightweight PDF font cannot represent the result."""


def _metadata_lines(result: Mapping[str, Any]) -> list[tuple[str, str]]:
    metadata = result.get("metadata") or {}
    return [
        ("Language", str(result.get("language", "N/A"))),
        ("Source word count", str(result.get("word_count", "N/A"))),
        ("Mode", str(metadata.get("mode", "N/A"))),
        ("Provider", str(metadata.get("provider", "N/A"))),
        ("Requested model", str(metadata.get("requested_model", "N/A"))),
        ("Routed model", str(metadata.get("routed_model") or "Not applicable")),
        ("Input format", str(metadata.get("input_format", "N/A")).upper()),
        ("Chunk count", str(metadata.get("chunk_count", "N/A"))),
        ("Processing time", f"{metadata.get('processing_time_ms', 'N/A')} ms"),
    ]


def build_markdown(result: Mapping[str, Any]) -> bytes:
    """Create a UTF-8 Markdown report containing the complete result contract."""

    points = "\n".join(f"- {point}" for point in result.get("key_points", []))
    details = "\n".join(f"- **{label}:** {value}" for label, value in _metadata_lines(result))
    text = (
        "# Document summary\n\n"
        f"{result.get('summary', '')}\n\n"
        "## Key points\n\n"
        f"{points}\n\n"
        "## Processing details\n\n"
        f"{details}\n"
    )
    return text.encode("utf-8")


def build_docx(result: Mapping[str, Any]) -> bytes:
    """Create a Unicode-safe Word report."""

    document = Document()
    document.add_heading("Document summary", level=1)
    document.add_paragraph(str(result.get("summary", "")))
    document.add_heading("Key points", level=2)
    for point in result.get("key_points", []):
        document.add_paragraph(str(point), style="List Bullet")
    document.add_heading("Processing details", level=2)
    for label, value in _metadata_lines(result):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(result: Mapping[str, Any]) -> bytes:
    """Create a Latin-1 PDF, rejecting rather than corrupting Unicode text."""

    summary = str(result.get("summary", ""))
    points = [str(point) for point in result.get("key_points", [])]
    lines = [summary, *points, *(f"{label}: {value}" for label, value in _metadata_lines(result))]
    try:
        "\n".join(lines).encode("latin-1")
    except UnicodeEncodeError as exc:
        raise PdfCharacterError(
            "PDF export supports Latin-1 text only; use Markdown or DOCX for Unicode."
        ) from exc

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(text="Document summary", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(w=0, text=summary, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(text="Key points", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    for point in points:
        pdf.multi_cell(w=0, text=f"- {point}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(text="Processing details", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    for label, value in _metadata_lines(result):
        pdf.multi_cell(
            w=0,
            text=f"{label}: {value}",
            new_x="LMARGIN",
            new_y="NEXT",
        )
    return bytes(pdf.output())
